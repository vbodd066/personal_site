from docx import Document

try:
    doc = Document('src/data/Harvard Amgen Resume 2026.docx')
    print("Resume Content:")
    print("=" * 80)
    
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    if doc.tables:
        print("\nTables Found:")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        print(cell.text)
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
